"""青志协加分数据库 — Excel → Word 生成脚本

用法:
    python generate_list.py 3.29

功能:
    从 Excel 年级花名册中读取指定日期的加分数据，
    按排版规则生成格式化的 Word 加分名单文档。
"""

import sys
import os
import re
from collections import defaultdict

sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import openpyxl
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils import (
    load_class_names, normalize_name, class_sort_key, name_length_key,
    classify_char, get_excel_path, get_output_dir,
    find_date_columns, extract_date_prefix
)


# ============================================================
# 字体常量
# ============================================================

FONT_SONG = '宋体'
FONT_TNR = 'Times New Roman'
FONT_SIZE = Pt(14)  # 四号

# 悬挂缩进：续行缩进 7.95 字符（对齐班级名+两个空格后）
# 对应模板格式：w:leftChars="795" w:left="1749"


# ============================================================
# 数据读取
# ============================================================

def read_scores_for_date(excel_path, date_str):
    """从 Excel 中读取指定日期列的所有加分数据。

    date_str 可以是纯日期如 '3.29'，也可以是完整列头如 '3.29清瓶乐'。
    使用 startswith 匹配，支持列头含活动名（如 '3.29清瓶乐'）。

    Returns:
        list of (class_name, name, score) 三元组，score 非空
    """
    wb = openpyxl.load_workbook(excel_path)
    class_names = load_class_names(wb)
    records = []

    for cn in sorted(class_names):
        ws = wb[cn]

        # 查找日期列（startswith 匹配
        date_col = None
        for col_idx in range(3, ws.max_column + 1):
            cell_value = ws.cell(row=1, column=col_idx).value
            if cell_value and str(cell_value).strip().startswith(date_str):
                date_col = col_idx
                break

        if date_col is None:
            continue

        # 读取该列所有非空分数
        for row_idx in range(2, ws.max_row + 1):
            name_cell = ws.cell(row=row_idx, column=2).value
            score_cell = ws.cell(row=row_idx, column=date_col).value

            if name_cell and score_cell is not None:
                name = str(name_cell).strip()
                try:
                    score = float(score_cell)
                except (ValueError, TypeError):
                    continue
                records.append((cn, name, score))

    wb.close()
    return records


# ============================================================
# 排序
# ============================================================

def sort_records(records):
    """按规则排序：分数降序 → 班级优先级 → 姓名长度。"""
    return sorted(
        records,
        key=lambda r: (
            -r[2],                     # 分数降序
            class_sort_key(r[0]),      # 班级优先级
            name_length_key(r[1]),     # 姓名长度
        )
    )


# ============================================================
# 姓名格式化
# ============================================================

def format_name_for_output(name):
    """将 Excel B 列的姓名格式化为 Word 输出格式。

    两字姓名：两字中间用两个半角空格（U+0020）填充对齐
    其他姓名：保持原样（去掉多余空格）
    """
    clean = normalize_name(name)
    if len(clean) == 2 and '·' not in name and '·' not in clean:
        return clean[0] + '  ' + clean[1]
    return name.strip()


# ============================================================
# Word 文档生成
# ============================================================

def add_mixed_font_run(paragraph, text, bold=False, no_wrap=False):
    """向段落添加文本，自动按字符类型分 run 设置字体。

    中文字符 → 宋体（通过 w:eastAsia）
    数字/括号/点(·) → Times New Roman（通过 w:ascii/w:hAnsi）
    空格 → 宋体（独立成组，不跟随相邻字符）
    no_wrap — 整个文本合并为单个 run 并追加 <w:noWrap/>，
              西文字体设为 Times New Roman，东亚字体设为宋体。
    """
    if not text:
        return

    # no_wrap 模式：完整字符串作为一个 run，不按字符类型拆分
    # 关键：西文字体设 Times New Roman，东亚字体设宋体
    if no_wrap:
        # 不移除 U+FEFF — 改用语言标记来触发 noLineBreaksAfter 规则
        run = paragraph.add_run(text)
        run.font.size = FONT_SIZE
        run.font.name = FONT_TNR
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)
        if bold:
            run.font.bold = True

        # 显式标记语言为中文，确保 w:noLineBreaksAfter 规则正确匹配
        rPr = run._element.get_or_add_rPr()
        lang_elem = OxmlElement('w:lang')
        lang_elem.set(qn('w:val'), 'zh-CN')
        lang_elem.set(qn('w:eastAsia'), 'zh-CN')
        rPr.append(lang_elem)

        # 含空格时加 preserve 保护
        if ' ' in text:
            t_elem = run._element.find(qn('w:t'))
            if t_elem is not None:
                t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        rPr.append(OxmlElement('w:noWrap'))
        return

    # 以下为非 no_wrap 模式（正常字符类型分组）
    groups = []
    current_type = None
    current_chars = []

    for ch in text:
        ctype = classify_char(ch)
        if ctype == 'space':
            if current_chars:
                groups.append((current_type, ''.join(current_chars)))
                current_type = None
                current_chars = []
            groups.append(('space', ch))
            continue
        if ctype != current_type:
            if current_chars:
                groups.append((current_type, ''.join(current_chars)))
            current_type = ctype
            current_chars = [ch]
        else:
            current_chars.append(ch)

    if current_chars:
        groups.append((current_type, ''.join(current_chars)))

    # 合并连续空格组，减少 run 数量
    merged = []
    for ctype, chunk in groups:
        if ctype == 'space' and merged and merged[-1][0] == 'space':
            merged[-1] = ('space', merged[-1][1] + chunk)
        else:
            merged.append((ctype, chunk))

    for ctype, chunk in merged:
        if ctype == 'space':
            run = paragraph.add_run(chunk)
            run.font.name = FONT_SONG
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)
            run.font.size = FONT_SIZE
            if bold:
                run.font.bold = True
            continue

        run = paragraph.add_run(chunk)
        run.font.size = FONT_SIZE

        if ctype == 'chinese':
            run.font.name = FONT_SONG
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)
        elif ctype == 'digit_punct':
            run.font.name = FONT_TNR
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)
        else:
            run.font.name = FONT_TNR
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)

        if bold:
            run.font.bold = True


def apply_hanging_indent(paragraph):
    """为段落设置悬挂缩进（6.25字符/1750 twips）。

    规则：首行缩进 6.25 字符（悬挂缩进），
    效果：第一行顶格写，续行缩进 6.25 字符。
    """
    pPr = paragraph._element.get_or_add_pPr()
    for existing in pPr.findall(qn('w:ind')):
        pPr.remove(existing)
    ind = OxmlElement('w:ind')
    # leftChars="625" = 6.25字符（以100为单位），left twips 备用值
    ind.set(qn('w:left'), '1750')
    ind.set(qn('w:hanging'), '1750')
    pPr.append(ind)


# ============================================================
# 阶段1：手动分行（绕过 Word CJK 断行引擎）
# ============================================================

# 页面参数
PAGE_WIDTH_CM = 21.0
MARGIN_LEFT_CM = 3.17
MARGIN_RIGHT_CM = 3.17
USABLE_WIDTH_CM = PAGE_WIDTH_CM - MARGIN_LEFT_CM - MARGIN_RIGHT_CM  # 14.66

PF_LEFT_CM = MARGIN_LEFT_CM
PF_RIGHT_CM = MARGIN_RIGHT_CM
PF_TOP_CM = 2.54
PF_BOTTOM_CM = 2.54

# 缩进：6.25 字符（排版规则）
INDENT_CHARS = 6.25

# 字符宽度（pt，对应四号字 14pt）
CJK_WIDTH = 14       # 中文全角字符宽度
ASCII_WIDTH = 14     # 数字/括号/中点（Times New Roman）宽度同中文全角
SPACE_WIDTH = 7      # 半角空格（只有半角字符宽度的一半）


def char_width_pt(ch):
    """返回单个字符的宽度（pt）。基于 14pt 四号字。"""
    if ch == ' ':
        return SPACE_WIDTH
    cp = ord(ch)
    # CJK 范围：全角宽度
    if (0x4E00 <= cp <= 0x9FFF or 0x3400 <= cp <= 0x4DBF or
        0x3000 <= cp <= 0x303F or 0xFF00 <= cp <= 0xFFEF or
        0x20000 <= cp <= 0x2FFFF or 0xF900 <= cp <= 0xFAFF or
        0x2F800 <= cp <= 0x2FA1F):
        return CJK_WIDTH
    # ASCII 范围：数字、拉丁字母、标点等
    return ASCII_WIDTH


def text_width_pt(text):
    """返回文本的总宽度（pt）。"""
    return sum(char_width_pt(ch) for ch in text)


def calc_lines_for_class(class_name, name_units):
    """计算一个班级的文本行分割。

    返回: [(text, is_first), ...]
    """
    usable_pt = (USABLE_WIDTH_CM / 2.54) * 72
    cont_pt = usable_pt - INDENT_CHARS * CJK_WIDTH

    lines = []

    # 先构建班级首行
    cls_text = class_name + '  '
    cls_width = text_width_pt(cls_text)
    first_line = cls_text
    first_width = cls_width

    # 收集续行上的姓名
    cont_units = []

    for unit in name_units:
        uw = text_width_pt(unit)
        if first_width + uw <= usable_pt:
            # 放首行
            first_line += unit
            first_width += uw
        else:
            # 放续行
            cont_units.append(unit)

    # 首行
    lines.append((first_line, True))

    # 续行：贪心填充
    if cont_units:
        current_line = ''
        current_width = 0.0
        for unit in cont_units:
            uw = text_width_pt(unit)
            if current_width + uw <= cont_pt:
                current_line += unit
                current_width += uw
            else:
                if current_line:
                    lines.append((current_line, False))
                current_line = unit
                current_width = uw
        if current_line:
            lines.append((current_line, False))

    # 后优化：如果两条同类行的内容可合并为一行，合并以减少总行数。
    # 必须使用目标行的宽度上限做溢出检查：续行用 cont_pt，首行用 usable_pt。
    # 只合并同类行（CONT+CONT 或 FIRST+FIRST），避免续行内容失去缩进。
    if len(lines) >= 2:
        last_text, last_is_first = lines[-1]
        prev_text, prev_is_first = lines[-2]
        if last_is_first == prev_is_first:
            last_w = text_width_pt(last_text)
            prev_w = text_width_pt(prev_text)
            limit = usable_pt if prev_is_first else cont_pt
            if prev_w + last_w <= limit:
                lines[-2] = (prev_text + last_text, prev_is_first)
                lines.pop()

    return lines


def build_name_unit(name):
    """构建姓名单元文本（含末尾跟随的两个分格空格）。

    两字姓名在两字中间插入两个半角空格。
    阶段1用占位符标记：中点 · → 用全角空格 U+3000 替代。
    """
    clean = normalize_name(name)
    # 两字姓名：中间加两个半角空格
    if len(clean) == 2 and '·' not in name:
        inner = clean[0] + '  ' + clean[1]
    else:
        inner = name.strip()
    # 中点替换为全角空格（跟其他字符统一宽度）
    inner = inner.replace('·', '　')
    # 末尾跟两个分格空格
    return inner + '  '


def wrap_lines_for_class(score_label, class_name, name_units):
    """将一个班级的全部姓名手动分行。

    班级首行：班级名 + 两个空格 + 姓名列表（顶格无缩进）
    续行：  无班级名，姓名列表（悬挂缩进 6.25 字符）

    返回: [(text, is_first), ...]
    """
    lines = []

    # 可用宽度 pt
    usable_pt = (USABLE_WIDTH_CM / 2.54) * 72
    cont_pt = usable_pt - INDENT_CHARS * CJK_WIDTH

    # 构建班级首行
    cls_text = class_name + '  '
    cls_width = text_width_pt(cls_text)

    # 先计算每个 name_unit 的宽度，如果首行放不下则提前创建续行
    # 策略：贪心填充，从后向前合并以减少续行数量
    first_names = []
    first_width = cls_width
    cont_names = []  # 存放 (unit_text, unit_width) 列表的列表
    current_cont = []
    current_cont_width = 0.0
    has_first_overflow = False

    for unit in name_units:
        uw = text_width_pt(unit)
        if not has_first_overflow and first_width + uw <= usable_pt:
            first_names.append(unit)
            first_width += uw
        else:
            has_first_overflow = True
            # 检查是否可以塞进上一条续行
            if current_cont_width + uw <= cont_pt:
                current_cont.append(unit)
                current_cont_width += uw
            else:
                if current_cont:
                    cont_names.append(current_cont)
                current_cont = [unit]
                current_cont_width = uw
    if current_cont:
        cont_names.append(current_cont)

    # 从后往前合并：如果倒数第2条续行有多余空间，试着从最后一条挪一个名字过来
    if len(cont_names) >= 2:
        for ci in range(len(cont_names)-2, -1, -1):
            while True:
                last = cont_names[ci+1]
                cur = cont_names[ci]
                if not last: break
                first_unit = last[0]
                fw = text_width_pt(first_unit)
                cur_w = sum(text_width_pt(u) for u in cur)
                if cur_w + fw <= cont_pt:
                    cur.append(first_unit)
                    last.pop(0)
                    if not last:
                        cont_names.pop(ci+1)
                else:
                    break

    # 构建输出行
    lines = []
    # 首行
    first_line = cls_text + ''.join(first_names)
    lines.append((first_line, True))

    # 续行
    for cn in cont_names:
        line = ''.join(cn)
        lines.append((line, False))

    return lines


def generate_all_lines(records):
    """将所有记录转为行列表，完成阶段1。

    返回: [(text, line_type), ...]
        line_type: 'title' | 'first' | 'cont'
    """
    by_score = defaultdict(list)
    for class_name, name, score in records:
        by_score[score].append((class_name, name))

    all_lines = []

    for score in sorted(by_score.keys(), reverse=True):
        entries = by_score[score]
        score_text = f"({int(score) if score == int(score) else score}分)"

        by_class = []
        cur_cls = None
        cur_names = []
        for class_name, name in entries:
            if class_name != cur_cls:
                if cur_cls is not None:
                    by_class.append((cur_cls, cur_names))
                cur_cls = class_name
                cur_names = [name]
            else:
                cur_names.append(name)
        if cur_cls is not None:
            by_class.append((cur_cls, cur_names))

        first_class = True
        for class_name, names in by_class:
            name_units = [build_name_unit(n) for n in names]
            score_label = score_text if first_class else None
            if score_label:
                all_lines.append((score_label, 'title'))
            # 使用 calc_lines_for_class 进行分行
            lines_out = calc_lines_for_class(class_name, name_units)
            for text, is_first in lines_out:
                all_lines.append((text, 'first' if is_first else 'cont'))
            first_class = False

    return all_lines


# ============================================================
# 阶段2：还原占位符
# ============================================================

def restore_placeholders(lines):
    """还原占位符：全角空格 → 中点。"""
    restored = []
    for text, ltype in lines:
        t = text.replace('　', '·')
        restored.append((t, ltype))
    return restored


# ============================================================
# 阶段3：按行写入 docx
# ============================================================

def add_mixed_font_run(paragraph, text, bold=False):
    """向段落按字符类型分 run 添加文本并设置字体。

    中文 → 宋体（w:eastAsia）
    数字/括号/中点· → Times New Roman（w:ascii/w:hAnsi）
    空格 → 宋体
    """
    if not text:
        return

    groups = []
    current_type = None
    current_chars = []

    for ch in text:
        ctype = classify_char(ch)
        if ctype == 'space':
            if current_chars:
                groups.append((current_type, ''.join(current_chars)))
                current_type = None
                current_chars = []
            groups.append(('space', ch))
            continue
        if ctype != current_type:
            if current_chars:
                groups.append((current_type, ''.join(current_chars)))
            current_type = ctype
            current_chars = [ch]
        else:
            current_chars.append(ch)

    if current_chars:
        groups.append((current_type, ''.join(current_chars)))

    # 合并连续空格
    merged = []
    for ctype, chunk in groups:
        if ctype == 'space' and merged and merged[-1][0] == 'space':
            merged[-1] = ('space', merged[-1][1] + chunk)
        else:
            merged.append((ctype, chunk))

    for ctype, chunk in merged:
        if ctype == 'space':
            run = paragraph.add_run(chunk)
            run.font.name = FONT_SONG
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)
            run.font.size = FONT_SIZE
            if bold:
                run.font.bold = True
            # 空格需要 preserve
            t_elem = run._element.find(qn('w:t'))
            if t_elem is not None:
                t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            continue

        run = paragraph.add_run(chunk)
        run.font.size = FONT_SIZE

        if ctype == 'chinese':
            run.font.name = FONT_SONG
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)
        elif ctype == 'digit_punct':
            run.font.name = FONT_TNR
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)
        else:
            run.font.name = FONT_TNR
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)

        if bold:
            run.font.bold = True


def apply_continuation_indent(paragraph):
    """设置续行左缩进（匹配排版规则 6.25 字符）。

    使用 leftChars 字符单位缩进 6.25 字符，
    缩进值 = 625（除以100为6.25字符）；
    twips 备用值 = 1750（6.25字符×14pt×20twips/pt）。
    """
    pPr = paragraph._element.get_or_add_pPr()
    for existing in pPr.findall(qn('w:ind')):
        pPr.remove(existing)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:leftChars'), '625')
    ind.set(qn('w:left'), '1750')
    pPr.append(ind)


def add_title_paragraph(par, text):
    """添加标题段落：(分数) 加粗，无缩进。"""
    add_mixed_font_run(par, text, bold=True)


def add_first_paragraph(par, text):
    """添加班级首行：顶格，无缩进。"""
    add_mixed_font_run(par, text)


def add_continuation_paragraph(par, text):
    """添加续行：悬挂缩进 6.25 字符。"""
    add_mixed_font_run(par, text)
    apply_continuation_indent(par)


def generate_word(records, date_str, output_path):
    """阶段1-4合一：手动分行 → 还原占位符 → 写入 docx（含行尾去空格） → 验证清理。"""
    doc = Document()

    # 兼容性设置
    settings = doc.settings.element
    for tag, attr in [('w:useFELayout', None),
                      ('w:balanceSingleByteDoubleByteWidth', None)]:
        if settings.find(qn(tag)) is None:
            el = OxmlElement(tag)
            settings.append(el)
    if settings.find(qn('w:characterSpacingControl')) is None:
        csc = OxmlElement('w:characterSpacingControl')
        csc.set(qn('w:val'), 'compressPunctuation')
        settings.append(csc)

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 默认字体
    style = doc.styles['Normal']
    style.font.name = FONT_SONG
    style.font.size = FONT_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0

    # ============ 阶段1：手动分行 ============
    all_lines = generate_all_lines(records)

    # ============ 阶段2：还原占位符 ============
    lines = restore_placeholders(all_lines)

    # ============ 阶段3：按行写入 docx（含行尾空格剥离） ============
    for text, ltype in lines:
        text = text.rstrip()  # 排版规则要求行尾不得有空格
        p = doc.add_paragraph()
        if ltype == 'title':
            add_title_paragraph(p, text)
        elif ltype == 'first':
            add_first_paragraph(p, text)
        elif ltype == 'cont':
            add_continuation_paragraph(p, text)

    # 保存
    doc.save(output_path)

    # ============ 阶段4：后处理验证 ============
    import zipfile
    with zipfile.ZipFile(output_path, 'r') as z:
        xml = z.read('word/document.xml').decode('utf-8')
    if '〈' in xml or '〉' in xml:
        print("警告: 残留书名号占位符")
    if '　' in xml:
        print("警告: 残留全角空格占位符")
        # 阶段3 已对每条输出行执行 text.rstrip()，行尾半角空格在写入前被剥离。
        # 如此处仍残留全角空格，说明 restore_placeholders 未完全还原，属于代码缺陷。

    total_names = len(records)
    dot_count = xml.count('·')
    expected_dots = sum(1 for _, name, _ in records if '·' in name)
    if dot_count != expected_dots:
        print(f"警告: 中点数量不匹配 (实际 {dot_count}, 期望 {expected_dots})")

    return output_path


# ============================================================
# 废弃函数（保留避免导入错误，不再使用）
# ============================================================

def format_name_for_output(name):
    """废弃：改用 build_name_unit。"""
    clean = normalize_name(name)
    if len(clean) == 2 and '·' not in name and '·' not in clean:
        return clean[0] + '  ' + clean[1]
    return name.strip()

def apply_hanging_indent(paragraph):
    """废弃：改用 apply_continuation_indent。"""
    apply_continuation_indent(paragraph)


# ============================================================
# 主流程
# ============================================================

def main():
    args = sys.argv[1:]

    # 默认优先使用 import_scores 生成的 updated 副本
    updated_path = os.path.join(get_output_dir(), '年级花名册数据表_updated.xlsx')
    if os.path.exists(updated_path):
        excel_path = updated_path
    else:
        excel_path = get_excel_path()

    # 解析 --excel 选项
    if '--excel' in args:
        idx = args.index('--excel')
        if idx + 1 < len(args):
            excel_path = args[idx + 1]
            args.pop(idx)
            args.pop(idx)
        else:
            print("错误: --excel 需要指定文件路径")
            sys.exit(1)

    if len(args) < 1:
        print("用法: python generate_list.py <date> [--excel <path>]")
        print("示例: python generate_list.py 3.29")
        sys.exit(1)

    date_str = args[0]

    print("=" * 60)
    print("青志协加分数据库 — 生成脚本 (Excel → Word)")
    print("=" * 60)
    print()

    # 检测同日期多活动
    from utils import scan_excel_date_columns
    columns = scan_excel_date_columns(excel_path, date_str)
    if len(columns) == 0:
        print(f"\n错误: 未找到日期 '{date_str}' 的加分数据。")
        sys.exit(1)
    elif len(columns) > 1:
        print(f"\n错误: 日期 '{date_str}' 对应多个活动列，请指定完整列头：")
        for c in columns:
            print(f"  {c}")
        print()
        sys.exit(1)

    # 读取数据
    print(f"数据库:   {os.path.basename(excel_path)}")
    print(f"活动日期: {date_str}")

    records = read_scores_for_date(excel_path, date_str)

    if not records:
        print(f"\n错误: 未找到日期 '{date_str}' 的加分数据。")
        sys.exit(1)

    print(f"读取记录: {len(records)} 条")
    print()

    # 排序
    records = sort_records(records)

    # 汇总
    score_dist = defaultdict(int)
    for _, _, score in records:
        score_dist[score] += 1
    print("分数分布:")
    for s in sorted(score_dist.keys(), reverse=True):
        print(f"  {s}分: {score_dist[s]} 人")
    print()

    # 生成 Word
    date_prefix = extract_date_prefix(date_str)
    output_filename = f"{date_prefix}清瓶乐活动德育分统计.docx"
    output_path = os.path.join(get_output_dir(), output_filename)

    generate_word(records, date_str, output_path)

    print(f"输出文件: {output_path}")
    print()
    print("生成完成！")


if __name__ == '__main__':
    main()
